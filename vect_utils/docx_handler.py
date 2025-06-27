import docx
from langchain_core.documents import Document

def handle_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception as e:
        print(f"❌ Error reading DOCX: {file_path}: {e}")
        return []
